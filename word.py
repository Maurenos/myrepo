 #!/usr/bin/python
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

document = Document()

paragraph = document.add_paragraph()
paragraph_format = paragraph.paragraph_format

p = document.add_paragraph('Name Surname ')

p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p = document.add_paragraph('Street Address ')

p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p = document.add_paragraph('ZIP code and City ')

p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

run = document.add_paragraph().add_run()
font = run.font
font.bold = True
z = document.add_paragraph('The Company to be applied to ')
z = document.add_paragraph('1st address line')
z = document.add_paragraph('2nd address line')

first = 'Your offer attracted my attention due to the following reasons.\
 First, XYZ is a corporation that offers broad development opportunities.\
 Secondly, the company provides many challenges and opportunities for its employees.\
 Thirdly, the company has a wide variety of development programs'
u = document.add_paragraph(first)
u.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

second = 'Taking into account my previous working experience and career interests,\
 I firmly believe that I would be an asset for your company that would increase your profitability.'

ux = document.add_paragraph(second)
ux.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

third = 'If you require additional information, do not hesitate to contact me.'

uxz = document.add_paragraph(third)
uxz.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

fourth = 'Regards,\nRadoslaw Nitka'

xyz = document.add_paragraph(fourth)
xyz.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

document.save('demo.docx')